import os
import docx
from docx.shared import Cm

#TODO: Format fonts, rounding, headings
#TODO: Load formatted Word template
#TODO: Add alternative export to csv or xlsx
#TODO: Subscript in table headings
#TODO: UserWarning: style lookup by style_id is deprecated. Use style name as key instead. return self._get_style_id_from_style(self[style_name], style_type)


class DataExport:
    def __init__(self):
        # Initialise the Word file
        self.doc = docx.Document()

    def spl_table(self, data, heading=None, decimals=0, dba_alignment="left"):
        #TODO: Format the column headers

        # Add the table heading
        assert heading is not None
        self.doc.add_heading(heading, 1)
        # Remove decimal points and 0
        data = data.round(decimals=decimals)
        #TODO: Move A-weighted columns to right

        # Initialise the table
        table = self.doc.add_table(rows=(data.shape[0] + 1), cols=data.shape[1] + 2, style="Table Grid")
        # table.style.TableGrid   # Add in borders
        # Add dates in first column
        table.cell(0, 0).text = "Position"  # Label first column
        table.cell(0, 1).text = "Date"
        ind = data.index.tolist()

        # Loop over the DataFrame and assign data to the Word Table
        # Position names
        for i in range(data.shape[0]):    # For each row
            current_cell = table.cell(i + 1, 0)
            prev_cell = table.cell(i, 0)
            if prev_cell.text == str(ind[i][0]):
                prev_cell.merge(current_cell)
            else:
                current_cell.text = str(ind[i][0])

        # Dates
            table.cell(i + 1, 1).text = str(ind[i][1])
            for j in range(data.shape[1]):  # Go through each column
                # Add column headings
                heading = str(data.columns[j][1])  # Remove index params from spectral column headings
                if len(data.columns) == 1:  # This occurs if dba_only=True
                    heading = str(data.columns[j])
                table.cell(0, j + 2).text = heading
                # And assign the values in the table.
                cell = data.iat[i, j]
                table.cell(i + 1, j + 2).text = str(cell)

    def weather_table(self, data, heading=None, decimals=1, metrics=["temp", "clouds", "wind_speed"]):
        assert heading is not None
        self.doc.add_heading(heading, 1)
        #TODO: Add a dict for metric units to be added to the table

        # Trim the rows
        data = data.loc[metrics]
        # Remove decimal points and 0
        data = data.round(decimals=decimals)
        # Initialise the table
        table = self.doc.add_table(rows=data.shape[0] + 1, cols=data.shape[1] + 1, style="Table Grid")
        table.cell(0, 0).text = "Metrics"
        # Add row names by looping through metrics and available data
        rows = data.index.to_list()
        for i in range(data.shape[0]):
            table.cell(i + 1, 0).text = str(rows[i]).capitalize()
        # Add column headings
        cols = data.columns.to_list()
        for i in range(data.shape[1]):
            table.cell(0, i + 1).text = str(cols[i]).capitalize()
        # Loop over the DataFrame and assign data to the word Table
        for i in range(data.shape[0]):  # For each row
            for j in range(data.shape[1]): # Go through each column
                # And assign the values in the table.
                cell = data.iat[i, j]
                table.cell(i + 1, j + 1).text = str(cell)

    def figure(self, img_pth, img_width=12):
        self.doc.add_picture(img_pth, width=Cm(img_width))
        #TODO: Add option for landscape images (requires section breaks)

    def export(self, path=None, filename="results.docx"):
        assert path is not None
        path = os.path.join(path, filename)
        self.doc.save(path)
